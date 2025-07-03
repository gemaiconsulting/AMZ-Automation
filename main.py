# -*- coding: utf-8 -*-
"""
AMZ-Automation: Fetch data from HubSpot, update Excel, and generate/upload NDAs, Proposals, SOWs, and MSAs to SharePoint.
"""
import json
import os
import time
import re
import io
import requests
import pandas as pd
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
import smtplib
from email.mime.text import MIMEText

# Load environment variables from .env
load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
# ENVIRONMENT VARIABLES
# ─────────────────────────────────────────────────────────────────────────────

# Azure / Microsoft Graph
CLIENT_ID            = os.getenv("CLIENT_ID")
CLIENT_SECRET        = os.getenv("CLIENT_SECRET")
TENANT_ID            = os.getenv("TENANT_ID")
SHAREPOINT_SITE_ID   = os.getenv("SHAREPOINT_SITE_ID")

# HubSpot
HUBSPOT_ACCESS_TOKEN = os.getenv("HUBSPOT_ACCESS_TOKEN")

# Asana
ASANA_TOKEN = os.getenv("ASANA_PERSONAL_ACCESS_TOKEN")
ASANA_TEAM_ID = os.getenv("ASANA_TEAM_ID")
ASANA_WORKSPACE_ID = os.getenv("ASANA_WORKSPACE_ID")

# OneDrive / Excel file IDs
CLIENT_DATA_FILE_ID = os.getenv("CLIENT_DATA_FILE_ID")  # ID of ClientData.xlsx
TEMPLATES_FOLDER_ID = os.getenv("TEMPLATES_FOLDER_ID")  # ID of '02. Internal'


# Subfolder IDs (under '02. Internal')
SUBFOLDER_01_NDA_ID           = os.getenv("SUBFOLDER_01_NDA_ID")
SUBFOLDER_02_PROPOSALS_ID     = os.getenv("SUBFOLDER_02_PROPOSALS_ID")
SUBFOLDER_03_CONTRACTS_ID     = os.getenv("SUBFOLDER_03_CONTRACTS_ID")
SUBFOLDER_04_SOWS_ID          = os.getenv("SUBFOLDER_04_SOWS_ID")
SUBFOLDER_05_MSAS_ID          = os.getenv("SUBFOLDER_05_MSAS_ID")

# NDA Templates
TEMPLATE_NDA_CANDIDATE_ID     = os.getenv("TEMPLATE_NDA_CANDIDATE_ID")
TEMPLATE_NDA_CONTRACTOR_ID    = os.getenv("TEMPLATE_NDA_CONTRACTOR_ID")
TEMPLATE_NDA_CORPORATE_ID     = os.getenv("TEMPLATE_NDA_CORPORATE_ID")

# Proposal Templates
PROPOSAL_TEMPLATE_RISK_ASSESSMENT_ID           = os.getenv("PROPOSAL_TEMPLATE_RISK_ASSESSMENT_ID")
PROPOSAL_TEMPLATE_CONSULTING_SERVICES_ID       = os.getenv("PROPOSAL_TEMPLATE_CONSULTING_SERVICES_ID")
PROPOSAL_TEMPLATE_RECRUITING_ID                = os.getenv("PROPOSAL_TEMPLATE_RECRUITING_ID")
PROPOSAL_TEMPLATE_TRAINING_ID                  = os.getenv("PROPOSAL_TEMPLATE_TRAINING_ID")
PROPOSAL_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID = os.getenv("PROPOSAL_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID")

# SOW Templates
SOW_TEMPLATE_RISK_ASSESSMENT_ID           = os.getenv("SOW_TEMPLATE_RISK_ASSESSMENT_ID")
SOW_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID = os.getenv("SOW_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID")
SOW_TEMPLATE_RECRUITING_ID                = os.getenv("SOW_TEMPLATE_RECRUITING_ID")
SOW_TEMPLATE_TRAINING_ID                  = os.getenv("SOW_TEMPLATE_TRAINING_ID")
SOW_TEMPLATE_CONSULTING_SERVICES_ID       = os.getenv("SOW_TEMPLATE_CONSULTING_SERVICES_ID")

# MSA Template
MSA_TEMPLATE_ID = os.getenv("MSA_TEMPLATE_ID")

# Remove hardcoded folder IDs and use environment variables
VENDORS_PARTNERS_FOLDER_ID = os.getenv("VENDORS_FOLDER_ID")
CLIENTS_FOLDER_ID = os.getenv("CLIENTS_FOLDER_ID")

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

GRAPH_API_BASE_URL  = "https://graph.microsoft.com/v1.0"
BASE_URL            = "https://api.hubapi.com/crm/v3/objects/"
PROPERTIES_API_URL  = "https://api.hubapi.com/properties/v1"
EXCEL_PATH          = "ClientData.xlsx"

# ─────────────────────────────────────────────────────────────────────────────
# AUTHENTICATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def get_ms_token():
    """
    Retrieve a new access token for Microsoft Graph API using client credentials.
    """
    token_url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
    payload = {
        "grant_type":    "client_credentials",
        "client_id":     CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "scope":         "https://graph.microsoft.com/.default"
    }
    response = requests.post(token_url, data=payload)
    token = response.json().get("access_token")
    if token:
        return token
    else:
        raise Exception(f"❌ Failed to authenticate with Microsoft Graph API: {response.json()}")

# Get a fresh MS Graph token and header
_access_token_ms = get_ms_token()
HEADERS_MS = {"Authorization": f"Bearer {_access_token_ms}"}

# HubSpot headers
HEADERS_HS = {
    "Authorization": f"Bearer {HUBSPOT_ACCESS_TOKEN}",
    "Content-Type":  "application/json"
}

# ─────────────────────────────────────────────────────────────────────────────
# UTILITY FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def clean_data(value):
    """
    Return 'N/A' if value is None or empty, else return the value.
    """
    return value if value else "N/A"

def send_error_email(subject, message):
    """
    Send an error email to itadmin@amzrisk.com via Outlook SMTP.
    Uses environment variables for sender email and password.
    Optionally allows SMTP server and port to be set via environment variables.
    """
    sender    = os.getenv("SMTP_SENDER")      # e.g. your_outlook_email@outlook.com
    recipient = "itadmin@amzrisk.com"
    password  = os.getenv("SMTP_PASSWORD")    # your Outlook app password
    smtp_server = os.getenv("SMTP_SERVER", "smtp.office365.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))

    msg = MIMEText(message)
    msg["Subject"] = subject
    msg["From"]    = sender
    msg["To"]      = recipient

    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender, password)
            server.sendmail(sender, recipient, msg.as_string())
    except Exception as e:
        print(f"❌ Failed to send error email: {e}")

def update_lead_source_for_contact(contact_id):
    """
    Set the contact's lead_source property to 'Website Contact Form'.
    """
    url = f"https://api.hubapi.com/crm/v3/objects/contacts/{contact_id}"
    payload = {"properties": {"lead_source": "Website Contact Form"}}
    resp = requests.patch(url, headers=HEADERS_HS, json=payload)
    if resp.status_code != 200:
        send_error_email("Lead Source Update Failed", resp.text)
    else:
        print(f"✅ Updated Lead Source for contact {contact_id}")

def update_lead_source_for_website_contacts(contacts):
    """
    For all contacts, if 'message' exists and is not empty, set lead_source to 'Website Contact Form'.
    """
    for c in contacts:
        contact_id = c.get("id")
        properties = c.get("properties", {})
        message = properties.get("message", "")
        lead_source = properties.get("lead_source", "")
        # Only update if message is non-empty and lead_source is not already set correctly
        if message and lead_source != "Website Contact Form":
            update_lead_source_for_contact(contact_id)


# ─────────────────────────────────────────────────────────────────────────────
# SYNC CLOSED-WON DEALS TO ASANA
# ─────────────────────────────────────────────────────────────────────────────

ASANA_HEADERS = {
    "Authorization": f"Bearer {ASANA_TOKEN}",
    "Content-Type": "application/json"
}

def get_existing_asana_projects(team_gid):
    url = f"https://app.asana.com/api/1.0/teams/{team_gid}/projects?archived=false"
    res = requests.get(url, headers=ASANA_HEADERS)
    if res.status_code == 200:
        return set(project["name"] for project in res.json().get("data", []))
    else:
        print("❌ Failed to fetch Asana projects:", res.text)
        return set()

def create_asana_project(project_name, workspace_gid, team_gid, existing_projects):
    if project_name in existing_projects:
        print(f"⏭️ Skipping '{project_name}' — already exists in Asana.")
        return
    url = "https://app.asana.com/api/1.0/projects"
    data = {
        "data": {
            "name": project_name,
            "workspace": workspace_gid,
            "team": team_gid
        }
    }
    res = requests.post(url, headers=ASANA_HEADERS, json=data)
    if res.status_code == 201:
        print(f"✅ Created Asana project: {project_name}")
    else:
        print(f"❌ Failed to create Asana project: {res.text}")

def sync_closed_won_deals_to_asana():
    print("🔄 Fetching Closed-Won deals from HubSpot...")
    url = "https://api.hubapi.com/crm/v3/objects/deals?properties=dealname,dealstage&limit=100"
    res = requests.get(url, headers=HEADERS_HS)
    if res.status_code != 200:
        print("❌ HubSpot API error:", res.text)
        return
    deals = res.json().get("results", [])
    existing_projects = get_existing_asana_projects(ASANA_TEAM_ID)
    for deal in deals:
        deal_name = deal["properties"].get("dealname", "Untitled Deal")
        deal_stage = deal["properties"].get("dealstage", "")
        if deal_stage == "contractsent":
            print(f"🎯 Deal '{deal_name}' is Closed-Won — checking Asana...")
            create_asana_project(deal_name, ASANA_WORKSPACE_ID, ASANA_TEAM_ID, existing_projects)




def iter_block_items(parent):
    """
    Yield each paragraph or table in a python-docx Document.
    """
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def replace_placeholder(paragraph, replacements):
    """
    Fully replace all placeholders in paragraph, even if broken across multiple runs.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    for key, val in replacements.items():
        full_text = full_text.replace(key, str(val) if val is not None else "")
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = full_text
        else:
            run.text = ""

def replace_placeholders_in_document(doc, replacements):
    """
    Apply replace_placeholder() to every paragraph and table cell in entire document
    including paragraphs, tables, headers, footers.
    """
    # Body paragraphs
    for p in doc.paragraphs:
        replace_placeholder(p, replacements)

    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder(p, replacements)

    # Headers
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_placeholder(p, replacements)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholder(p, replacements)

    # Footers
    for section in doc.sections:
        for p in section.footer.paragraphs:
            replace_placeholder(p, replacements)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholder(p, replacements)


# ─────────────────────────────────────────────────────────────────────────────
# HUBSPOT PROPERTY FETCHING
# ─────────────────────────────────────────────────────────────────────────────

def get_all_properties(object_type):
    """
    Fetch all available fields (properties) for a given HubSpot object type.
    """
    url = f"{PROPERTIES_API_URL}/{object_type}/properties"
    response = requests.get(url, headers=HEADERS_HS)
    if response.status_code == 200:
        properties = response.json()
        return [prop["name"] for prop in properties]
    else:
        print(f"❌ Failed to fetch properties for {object_type}: {response.json()}")
        return []

# Retrieve all fields dynamically
CONTACT_FIELDS = get_all_properties("contacts")
COMPANY_FIELDS = get_all_properties("companies")
DEAL_FIELDS   = get_all_properties("deals")

print(f"✅ Retrieved {len(CONTACT_FIELDS)} Contact fields")
print(f"✅ Retrieved {len(COMPANY_FIELDS)} Company fields")
print(f"✅ Retrieved {len(DEAL_FIELDS)} Deal fields")

def fetch_all_hubspot_data(object_type, fields):
    """
    Fetch all records from HubSpot API using pagination.
    """
    properties = ",".join(fields)
    all_records = []
    url = f"{BASE_URL}{object_type}?limit=100&properties={properties}"
    has_more = True

    while has_more:
        response = requests.get(url, headers=HEADERS_HS)
        if response.status_code == 200:
            data = response.json()
            all_records.extend(data.get("results", []))
            paging = data.get("paging", {})
            next_page = paging.get("next", {}).get("after")
            if next_page:
                url = f"{BASE_URL}{object_type}?limit=100&properties={properties}&after={next_page}"
            else:
                has_more = False
        else:
            print(f"❌ Failed to fetch {object_type}: {response.json()}")
            break

    return all_records

# ─────────────────────────────────────────────────────────────────────────────
# LOAD & UPDATE EXCEL (ClientData.xlsx)
# ─────────────────────────────────────────────────────────────────────────────

def load_existing_data(file_path):
    """
    Load existing ClientData.xlsx (all sheets). If not found, return None.
    """
    try:
        existing_data = pd.read_excel(file_path, sheet_name=None)
        print("📂 Existing data loaded.")
        return existing_data
    except FileNotFoundError:
        print("❌ No existing file found. Creating a new one.")
        return None

def update_or_append_data(existing_data, new_data, unique_field):
    """
    Merge new data with existing data:
    - If a record exists (matching unique_field), update it.
    - If a record is new, append it.
    """
    if existing_data is not None:
        existing_df = existing_data  # already a DataFrame, no need to convert
        if unique_field in existing_df.columns:
            merged_df = pd.concat([existing_df, new_data]).drop_duplicates(subset=[unique_field], keep="last")
        else:
            print(f"⚠️ Unique field '{unique_field}' not found in existing data. Appending all records.")
            merged_df = pd.concat([existing_df, new_data])
    else:
        merged_df = new_data

    return merged_df

# Fetch all HubSpot data
contacts = fetch_all_hubspot_data("contacts", CONTACT_FIELDS)
update_lead_source_for_website_contacts(contacts)   # <--- add here!

companies = fetch_all_hubspot_data("companies", COMPANY_FIELDS)
deals    = fetch_all_hubspot_data("deals", DEAL_FIELDS)

# Clean and convert to DataFrames
cleaned_contacts  = [{key: clean_data(c["properties"].get(key)) for key in CONTACT_FIELDS} for c in contacts]
cleaned_companies = [{key: clean_data(c["properties"].get(key)) for key in COMPANY_FIELDS} for c in companies]
cleaned_deals     = [{key: clean_data(d["properties"].get(key)) for key in DEAL_FIELDS} for d in deals]

contacts_df  = pd.DataFrame(cleaned_contacts)
companies_df = pd.DataFrame(cleaned_companies)
deals_df     = pd.DataFrame(cleaned_deals)

# Load and merge with existing Excel
existing_data = load_existing_data(EXCEL_PATH)
contacts_df  = update_or_append_data(existing_data.get("Contacts")  if existing_data else None, contacts_df,  "email")
companies_df = update_or_append_data(existing_data.get("Companies") if existing_data else None, companies_df, "name")
deals_df     = update_or_append_data(existing_data.get("Deals")     if existing_data else None, deals_df,     "dealname")

# Save updated Excel
with pd.ExcelWriter(EXCEL_PATH) as writer:
    contacts_df.to_excel(writer, sheet_name="Contacts", index=False)
    companies_df.to_excel(writer, sheet_name="Companies", index=False)
    deals_df.to_excel(writer, sheet_name="Deals", index=False)

print("✅ ClientData.xlsx updated: new and existing data merged.")

# ─────────────────────────────────────────────────────────────────────────────
# UPLOAD UPDATED EXCEL TO ONE DRIVE / SHAREPOINT
# ─────────────────────────────────────────────────────────────────────────────

def upload_file_to_onedrive(file_path):
    """
    Upload ClientData.xlsx to OneDrive under the SharePoint site.
    """
    file_name = os.path.basename(file_path)
    upload_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{CLIENTS_FOLDER_ID}:/{file_name}:/content"
    )
    with open(file_path, "rb") as f:
        response = requests.put(upload_url, headers=HEADERS_MS, data=f)
    if response.status_code in [200, 201]:
        print(f"✅ {file_name} uploaded successfully to OneDrive!")
    else:
        print(f"❌ Failed to upload {file_name}: {response.json()}")

upload_file_to_onedrive(EXCEL_PATH)

# ─────────────────────────────────────────────────────────────────────────────
# CREATE CLIENT FOLDERS & COPY SUBFOLDERS
# ─────────────────────────────────────────────────────────────────────────────

def download_company_data_sheet():
    """
    Download ClientData.xlsx from SharePoint and return the 'Companies' sheet as a DataFrame.
    """
    url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{CLIENT_DATA_FILE_ID}/content"
    )
    response = requests.get(url, headers=HEADERS_MS)
    if response.status_code == 200:
        file_stream = io.BytesIO(response.content)
        return pd.read_excel(file_stream, sheet_name="Companies")
    else:
        print(f"❌ Failed to download ClientData.xlsx: {response.json()}")
        return None

def sanitize_folder_name(name: str) -> str:
    """
    Sanitize folder name to be compatible with SharePoint:
      • Remove invalid characters: \ / : * ? " < > | # { } % ~ & ,
      • Collapse runs of whitespace into a single space,
      • Trim leading/trailing spaces,
      • Strip any trailing dot.
    """
    # 1) Remove "hard" illegal chars plus comma
    cleaned = re.sub(r'[\\/:*?"<>|#{}%~&,]', '', name)

    # 2) Collapse any run of whitespace into a single space
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()

    # 3) If it ends with a dot, strip it off
    cleaned = cleaned.rstrip('.')

    return cleaned

def get_or_create_company_folder(company_name, company_type):
    """
    Create a folder for the company in the correct parent folder based on type.
    Returns the folder ID and a flag indicating if subfolders should be created (True for clients, False for vendors/partners).
    """
    safe_name = sanitize_folder_name(company_name)
    if company_type.strip().lower() in ["vendor", "partner"]:
        parent_id = VENDORS_PARTNERS_FOLDER_ID
        allow_subfolders = False
    else:
        parent_id = CLIENTS_FOLDER_ID
        allow_subfolders = True
    url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{parent_id}/children"
    response = requests.get(url, headers=HEADERS_MS)
    if response.status_code == 200:
        folders = response.json().get("value", [])
        for folder in folders:
            if folder["name"] == safe_name:
                return folder["id"], allow_subfolders
    # Create new folder
    create_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{parent_id}/children"
    payload = {
        "name": safe_name,
        "folder": {},
        "@microsoft.graph.conflictBehavior": "fail"
    }
    create_resp = requests.post(create_url, headers=HEADERS_MS, json=payload)
    if create_resp.status_code == 201:
        return create_resp.json()["id"], allow_subfolders
    else:
        print(f"❌ Failed to create folder '{safe_name}': {create_resp.json()}")
        return None, allow_subfolders

# Update the main folder creation loop
companies_df = download_company_data_sheet()
if companies_df is None:
    raise Exception("❌ Unable to load company data!")

company_names = companies_df["name"].dropna().unique().tolist()
client_folders = {}

for idx, row in companies_df.iterrows():
    name = row.get("name")
    ctype = str(row.get("type", "")).strip()
    if not name:
        continue
    folder_id, allow_subfolders = get_or_create_company_folder(name, ctype)
    if folder_id:
        client_folders[name] = {"id": folder_id, "allow_subfolders": allow_subfolders}

print("📂 All company folders created:", client_folders)

# ─────────────────────────────────────────────────────────────────────────────
# NDA GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def fetch_primary_contact_nda(company_id):
    """
    Fetch the primary contact for a company, return properties including id, firstname, lastname, email, jobtitle, nda_status, contact_type.
    """
    assoc_url = f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}/associations/contacts"
    assoc_resp = requests.get(assoc_url, headers=HEADERS_HS)
    if assoc_resp.status_code != 200:
        return None
    for contact in assoc_resp.json().get("results", []):
        contact_id = contact.get("id")
        contact_url = (
            f"https://api.hubapi.com/crm/v3/objects/contacts/{contact_id}"
            "?properties=firstname,lastname,email,jobtitle,hs_lead_status,contact_type,nda_status"
        )
        contact_resp = requests.get(contact_url, headers=HEADERS_HS)
        if contact_resp.status_code == 200:
            props = contact_resp.json().get("properties", {})
            if props.get("firstname") and props.get("lastname"):
                props["id"] = contact_id
                return props
    return None

def update_contact_nda_status(contact_id):
    """
    Update a contact's nda_status property to 'Generated' in HubSpot.
    """
    url = f"https://api.hubapi.com/crm/v3/objects/contacts/{contact_id}"
    payload = {"properties": {"nda_status": "Generated"}}
    resp = requests.patch(url, headers=HEADERS_HS, json=payload)
    if resp.status_code != 200:
        send_error_email("Contact NDA Status Update Failed", resp.text)

def update_company_nda_status(company_id):
    """
    Update a company's nda_status property to 'generated' in HubSpot.
    """
    if not company_id:
        return
    url = f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}"
    payload = {"properties": {"nda_status": "generated"}}
    resp = requests.patch(url, headers=HEADERS_HS, json=payload)
    if resp.status_code != 200:
        send_error_email("Company NDA Status Update Failed", resp.text)

def get_or_create_subfolder(parent_folder_id, subfolder_name, template_folder_id):
    """
    Check if a subfolder exists in the parent folder; if not, create it by copying from template.
    Returns the subfolder ID.
    """
    # Check if subfolder exists
    url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{parent_folder_id}/children"
    response = requests.get(url, headers=HEADERS_MS)
    if response.status_code == 200:
        folders = response.json().get("value", [])
        for folder in folders:
            if folder["name"] == subfolder_name:
                return folder["id"]

    # Create new subfolder by copying from template
    copy_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{template_folder_id}/copy"
    payload = {
        "parentReference": {"id": parent_folder_id},
        "name": subfolder_name
    }
    copy_resp = requests.post(copy_url, headers=HEADERS_MS, json=payload)
    if copy_resp.status_code not in [200, 202]:
        print(f"❌ Failed to create subfolder '{subfolder_name}': {copy_resp.json()}")
        return None

    # Wait for copy to complete and get the new folder ID
    time.sleep(2)
    response = requests.get(url, headers=HEADERS_MS)
    if response.status_code == 200:
        folders = response.json().get("value", [])
        for folder in folders:
            if folder["name"] == subfolder_name:
                return folder["id"]
    
    return None

def generate_nda_for_company(company):
    """
    For each company (from HubSpot), generate and upload an NDA if needed.
    """
    company_id = company["id"]
    props = company["properties"]
    company_name = props.get("name", "Unknown Company")
    company_nda_status = (props.get("nda_status") or "").strip().lower()

    allow_subfolders = client_folders.get(company_name, {}).get("allow_subfolders", True)

    contact = fetch_primary_contact_nda(company_id)
    if not contact:
        return
    contact_nda_status = (contact.get("nda_status") or "").strip().lower()

    if contact_nda_status != "generate" and company_nda_status != "generate":
        return

    contact_type = (contact.get("contact_type") or "").strip().lower()
    if contact_type == "candidate":
        template_id = TEMPLATE_NDA_CANDIDATE_ID
    elif contact_type in ["contractor", "employee", "partner/producer"]:
        template_id = TEMPLATE_NDA_CONTRACTOR_ID
    else:
        template_id = TEMPLATE_NDA_CORPORATE_ID

    # Locate company folder in SharePoint
    parent_id = VENDORS_PARTNERS_FOLDER_ID if not allow_subfolders else CLIENTS_FOLDER_ID
    url_fldr = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{parent_id}/children"
    )
    folders = requests.get(url_fldr, headers=HEADERS_MS).json().get("value", [])
    company_folder = next((f for f in folders if f["name"] == company_name), None)
    if not company_folder:
        return

    # Determine target folder for NDA
    if allow_subfolders:
        # Create NDA subfolder if it doesn't exist
        nda_folder_id = get_or_create_subfolder(
            company_folder["id"],
            "01. NDA",
            SUBFOLDER_01_NDA_ID
        )
        if not nda_folder_id:
            return
        target_folder_id = nda_folder_id
    else:
        # Vendors/partners: use company folder directly
        target_folder_id = company_folder["id"]

    contact_name = f"{contact.get('firstname','').strip()}_{contact.get('lastname','').strip()}"
    filename = f"AMZ Risk - {contact_type.title()} NDA - {contact_name} - {datetime.now().strftime('%Y%m%d')}.docx"
    copy_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{template_id}/copy"
    payload = {"parentReference": {"id": target_folder_id}, "name": filename}
    copy_resp = requests.post(copy_url, headers=HEADERS_MS, json=payload)
    if copy_resp.status_code not in [200, 202]:
        send_error_email("NDA Copy Failed", copy_resp.text)
        return

    # Wait for copy to complete
    time.sleep(5)
    children_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{target_folder_id}/children"
    )
    items = requests.get(children_url, headers=HEADERS_MS).json().get("value", [])
    copied_file = next((f for f in items if f["name"] == filename), None)
    if not copied_file:
        send_error_email("NDA Not Found", f"Copy succeeded but file not found for {company_name}")
        return

    # Download, replace placeholders, re-upload
    download_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{copied_file['id']}/content"
    )
    file_data = requests.get(download_url, headers=HEADERS_MS).content
    with open(filename, "wb") as f:
        f.write(file_data)

    doc = Document(filename)
    placeholders = {
        "{legal_entity_name}": props.get("legal_entity_name", ""),
        "{address}":  props.get("address", ""),
        "{city}":     props.get("city", ""),
        "{state_list}": props.get("state_list", ""),
        "{zip}":      props.get("zip", ""),
        "{email}":    contact.get("email", ""),
        "{firstname}": contact.get("firstname", ""),
        "{lastname}":  contact.get("lastname", ""),
        "{jobtitle}":  contact.get("jobtitle", "")
    }
    replace_placeholders_in_document(doc, placeholders)
    doc.save(filename)

    # Upload filled NDA
    upload_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{target_folder_id}:/{filename}:/content"
    )
    with open(filename, "rb") as f:
        requests.put(upload_url, headers=HEADERS_MS, data=f)

    update_contact_nda_status(contact.get("id"))
    update_company_nda_status(company_id)
    print(f"✅ NDA '{filename}' generated and uploaded for {company_name}!")

# Run NDA generation
companies_list = fetch_all_hubspot_data("companies", COMPANY_FIELDS)
for comp in companies_list:
    generate_nda_for_company(comp)
print("✅ All NDAs processed!")

# ─────────────────────────────────────────────────────────────────────────────
# PROPOSAL GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def fetch_deals_for_proposal():
    """
    Fetch HubSpot deals with properties needed for proposal generation.
    """
    url = (
        "https://api.hubapi.com/crm/v3/objects/deals?properties="
        "dealname,proposal_status,proposal___service_line,hubspot_owner_id"
    )
    resp = requests.get(url, headers=HEADERS_HS)
    return resp.json().get("results", []) if resp.status_code == 200 else []

def fetch_associated_company_id_for_deal(deal_id):
    """
    Given a HubSpot deal ID, return the associated company ID.
    """
    url = f"https://api.hubapi.com/crm/v3/objects/deals/{deal_id}/associations/companies"
    resp = requests.get(url, headers=HEADERS_HS)
    results = resp.json().get("results", []) if resp.status_code == 200 else []
    return results[0]["id"] if results else None

def fetch_company_data_for_proposal(company_id):
    """
    Fetch company properties (name, city, state, zip, address) for proposal.
    """
    url = (
        f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}?properties="
        "name,city,state_list,zip,address"
    )
    resp = requests.get(url, headers=HEADERS_HS)
    return resp.json().get("properties", {}) if resp.status_code == 200 else {}

def fetch_primary_contact_for_proposal(company_id):
    """
    Fetch primary contact properties (firstname, lastname, email).
    """
    url = f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}/associations/contacts"
    resp = requests.get(url, headers=HEADERS_HS)
    results = resp.json().get("results", []) if resp.status_code == 200 else []
    if not results:
        return {}
    cid = results[0]["id"]
    contact_url = (
        f"https://api.hubapi.com/crm/v3/objects/contacts/{cid}?properties="
        "firstname,lastname,email"
    )
    contact_resp = requests.get(contact_url, headers=HEADERS_HS)
    return contact_resp.json().get("properties", {}) if contact_resp.status_code == 200 else {}

def fetch_owner_details(owner_id):
    """
    Fetch HubSpot owner first and last name, plus email.
    """
    url = f"https://api.hubapi.com/crm/v3/owners/{owner_id}"
    resp = requests.get(url, headers=HEADERS_HS)
    if resp.status_code == 200:
        data = resp.json()
        return f"{data.get('firstName','')} {data.get('lastName','')}".strip(), data.get("email", "")
    return "", ""

def update_proposal_status(deal_id):
    """
    Update a HubSpot deal's proposal_status to 'Generated'.
    """
    url = f"https://api.hubapi.com/crm/v3/objects/deals/{deal_id}"
    payload = {"properties": {"proposal_status": "Generated"}}
    resp = requests.patch(url, headers=HEADERS_HS, json=payload)
    if resp.status_code != 200:
        send_error_email("Proposal Status Update Failed", resp.text)

def proposal_exists_for_service_line(folder_id, company_name, service_line):
    """
    Check if any proposal for the same company and service line already exists,
    regardless of the date suffix.
    """
    prefix = f"AMZ Risk - {company_name} - Proposal - {service_line}"
    url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{folder_id}/children"
    resp = requests.get(url, headers=HEADERS_MS)
    return any(item["name"].startswith(prefix) for item in resp.json().get("value", []))

PROPOSAL_TEMPLATES = {
    "Risk Assessment":             PROPOSAL_TEMPLATE_RISK_ASSESSMENT_ID,
    "Consulting Services":         PROPOSAL_TEMPLATE_CONSULTING_SERVICES_ID,
    "Recruiting":                  PROPOSAL_TEMPLATE_RECRUITING_ID,
    "Training":                    PROPOSAL_TEMPLATE_TRAINING_ID,
    "Global Threat Intelligence":  PROPOSAL_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID
}

def generate_proposal_for_deal(deal):
    """
    For each deal, generate and upload a Proposal if proposal_status is 'generate'.
    """
    status = deal["properties"].get("proposal_status", "").strip().lower()
    if status != "generate":
        return

    deal_id = deal["id"]
    company_id = fetch_associated_company_id_for_deal(deal_id)
    if not company_id:
        return
    company = fetch_company_data_for_proposal(company_id)
    company_name = company.get("name", "Unknown Company")
    contact = fetch_primary_contact_for_proposal(company_id)


    # Check if subfolders are allowed for this company
    allow_subfolders = client_folders.get(company_name, {}).get("allow_subfolders", True)
    if not allow_subfolders:
        print(f"⏩ Skipping Proposal generation for vendor/partner: {company_name}")
        return

    raw = deal["properties"].get("proposal___service_line", [])
    service_lines = []
    if isinstance(raw, list):
        service_lines = [opt.get("value") for opt in raw if isinstance(opt, dict) and opt.get("value")]
    elif isinstance(raw, str) and raw:
        service_lines = [s.strip() for s in raw.split(';') if s.strip()]
    if not service_lines:
        service_lines = ["Risk Assessment"]

    owner_id = deal["properties"].get("hubspot_owner_id")
    owner_name, owner_email = fetch_owner_details(owner_id)

    # Locate client folder
    url_fldr = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{CLIENTS_FOLDER_ID}/children"
    )
    folders = requests.get(url_fldr, headers=HEADERS_MS).json().get("value", [])
    client_folder = next((f for f in folders if f["name"] == company_name), None)
    if not client_folder:
        return

    # Create Proposals subfolder if it doesn't exist
    proposals_folder_id = get_or_create_subfolder(
        client_folder["id"],
        "02. Proposals",
        SUBFOLDER_02_PROPOSALS_ID
    )
    if not proposals_folder_id:
        return

    children_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{proposals_folder_id}/children"
    )
    for service_line in service_lines:
        filename = (
            f"AMZ Risk - {company_name} - Proposal - {service_line} - "
            f"{datetime.now().strftime('%Y%m%d')}.docx"
        )
        if proposal_exists_for_service_line(proposals_folder_id, company_name, service_line):
            print(f"⏩ Skipping duplicate proposal for {company_name} - {service_line}")    
            continue

        template_id = PROPOSAL_TEMPLATES.get(service_line, PROPOSAL_TEMPLATES["Risk Assessment"])
        copy_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{template_id}/copy"
        copy_resp = requests.post(
            copy_url,
            headers=HEADERS_MS,
            json={"parentReference": {"id": proposals_folder_id}, "name": filename}
        )
        if copy_resp.status_code not in (200, 202):
            send_error_email("Proposal Copy Failed", copy_resp.text)
            continue

        # Wait for copy to complete
        for _ in range(10):
            items = requests.get(children_url, headers=HEADERS_MS).json().get("value", [])
            if any(item["name"] == filename for item in items):
                break
            time.sleep(2)
        else:
            send_error_email("Proposal Copy Timeout", f"Copy timed out for {filename}")
            continue

        copied = next(item for item in items if item["name"] == filename)
        download_url = (
            f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
            f"/drive/items/{copied['id']}/content"
        )
        data = requests.get(download_url, headers=HEADERS_MS).content
        with open(filename, "wb") as fd:
            fd.write(data)

        placeholders = {
            "{proposal___service_line}": service_line,
            "{today’s date}":            datetime.now().strftime("%Y-%m-%d"),
            "{firstname}":               contact.get("firstname", ""),
            "{lastname}":                contact.get("lastname", ""),
            "{email}":                   contact.get("email", ""),
            "{legal_entity_name}":       company.get("legal_entity_name", ""),
            "{address}":                 company.get("address", ""),
            "{city}":                    company.get("city", ""),
            "{state_list}":              company.get("state_list", ""),
            "{zip}":                     company.get("zip", ""),
            "{amz_rep}":                 owner_name,
            "{amz_rep_email}":           owner_email
        }

        doc = Document(filename)
        replace_placeholders_in_document(doc, placeholders)
        doc.save(filename)

        upload_url = (
            f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
            f"/drive/items/{proposals_folder_id}:/{filename}:/content"
        )
        with open(filename, "rb") as fd:
            requests.put(upload_url, headers=HEADERS_MS, data=fd)

        update_proposal_status(deal_id)
        print(f"✅ Proposal '{filename}' uploaded for {company_name}!")

# Run proposal generation
deals_list = fetch_deals_for_proposal()
for deal in deals_list:
    generate_proposal_for_deal(deal)
print("✅ All proposals processed!")

# ─────────────────────────────────────────────────────────────────────────────
# SOW GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def fetch_deals_for_sow():
    """
    Fetch HubSpot deals with properties needed for SOW generation.
    """
    url = (
        "https://api.hubapi.com/crm/v3/objects/deals?properties="
        "dealname,sow_status,proposal___service_line,hubspot_owner_id"
    )
    resp = requests.get(url, headers=HEADERS_HS)
    return resp.json().get("results", []) if resp.status_code == 200 else []

def update_sow_status(deal_id):
    """
    Update a HubSpot deal's sow_status to 'Generated'.
    """
    url = f"https://api.hubapi.com/crm/v3/objects/deals/{deal_id}"
    payload = {"properties": {"sow_status": "Generated"}}
    resp = requests.patch(url, headers=HEADERS_HS, json=payload)
    if resp.status_code != 200:
        send_error_email("SOW Status Update Failed", resp.text)

SOW_TEMPLATES = {
    "Risk Assessment":            SOW_TEMPLATE_RISK_ASSESSMENT_ID,
    "Global Threat Intelligence": SOW_TEMPLATE_GLOBAL_THREAT_INTELLIGENCE_ID,
    "Recruiting":                 SOW_TEMPLATE_RECRUITING_ID,
    "Training":                   SOW_TEMPLATE_TRAINING_ID,
    "Consulting Services":        SOW_TEMPLATE_CONSULTING_SERVICES_ID
}

def generate_sow_for_deal(deal):
    """
    For each deal, generate and upload a SOW if needed.
    """
    deal_id = deal["id"]
    company_id = fetch_associated_company_id_for_deal(deal_id)
    if not company_id:
        return
    company = fetch_company_data_for_proposal(company_id)
    company_name = company.get("name", "Unknown Company")

    contact = fetch_primary_contact_for_proposal(company_id)

    # Check if subfolders are allowed for this company
    allow_subfolders = client_folders.get(company_name, {}).get("allow_subfolders", True)
    if not allow_subfolders:
        print(f"⏩ Skipping SOW generation for vendor/partner: {company_name}")
        return

    raw = deal["properties"].get("proposal___service_line", [])
    service_lines = []
    if isinstance(raw, list):
        service_lines = [opt.get("value") for opt in raw if isinstance(opt, dict) and opt.get("value")]
    elif isinstance(raw, str) and raw:
        service_lines = [s.strip() for s in raw.split(';') if s.strip()]
    if not service_lines:
        service_lines = ["Risk Assessment"]

    owner_id = deal["properties"].get("hubspot_owner_id")
    owner_name, owner_email = fetch_owner_details(owner_id)

    # Locate client folder
    url_fldr = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{CLIENTS_FOLDER_ID}/children"
    )
    folders = requests.get(url_fldr, headers=HEADERS_MS).json().get("value", [])
    client_folder = next((f for f in folders if f["name"] == company_name), None)
    if not client_folder:
        return

    
    # Locate or create SOW subfolder on demand
    url_sub = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{client_folder['id']}/children"
    )
    subfolders = requests.get(url_sub, headers=HEADERS_MS).json().get("value", [])
    sow_folder = next((f for f in subfolders if f["name"] == "04. SOWs"), None)

    if not sow_folder:
        # Create '04. SOWs' subfolder only if it doesn't exist
        create_folder_url = (
            f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{client_folder['id']}/children"
        )
        folder_payload = {
            "name": "04. SOWs",
            "folder": {},
            "@microsoft.graph.conflictBehavior": "fail"
        }
        create_resp = requests.post(create_folder_url, headers=HEADERS_MS, json=folder_payload)
        if create_resp.status_code in (200, 201):
            sow_folder = create_resp.json()
        else:
            send_error_email("SOW Subfolder Creation Failed", create_resp.text)
            return

    children_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{sow_folder['id']}/children"
    )
    for service_line in service_lines:
        filename = (
            f"AMZ Risk - {company_name} - SOW - {service_line} - "
            f"{datetime.now().strftime('%Y%m%d')}.docx"
        )
        if any(item["name"] == filename for item in requests.get(children_url, headers=HEADERS_MS).json().get("value", [])):
            continue

        template_id = SOW_TEMPLATES.get(service_line)
        if not template_id:
            continue

        copy_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{template_id}/copy"
        resp = requests.post(
            copy_url,
            headers=HEADERS_MS,
            json={"parentReference": {"id": sow_folder['id']}, "name": filename}
        )
        if resp.status_code not in (200, 202):
            send_error_email("SOW Copy Failed", resp.json())
            continue

        # Wait for copy
        for _ in range(10):
            items = requests.get(children_url, headers=HEADERS_MS).json().get("value", [])
            if any(i["name"] == filename for i in items):
                break
            time.sleep(2)
        else:
            send_error_email("SOW Copy Timeout", f"Copy timed out for {filename}")
            continue

        copied = next(i for i in items if i["name"] == filename)
        download_url = (
            f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
            f"/drive/items/{copied['id']}/content"
        )
        data = requests.get(download_url, headers=HEADERS_MS).content
        with open(filename, "wb") as fd:
            fd.write(data)

        placeholders = {
            "{proposal___service_line}": service_line,
            "{today's date}":            datetime.now().strftime("%Y-%m-%d"),
            "{firstname}":               contact.get("firstname", ""),
            "{lastname}":                contact.get("lastname", ""),
            "{jobtitle}":                contact.get("jobtitle", ""),
            "{email}":                   contact.get("email", ""),
            "{legal_entity_name}":       company.get("legal_entity_name", ""),
            "{address}":                 company.get("address", ""),
            "{city}":                    company.get("city", ""),
            "{state_list}":              company.get("state_list", ""),
            "{zip}":                     company.get("zip", ""),
            "{amz_rep}":                 owner_name,
            "{amz_rep_email}":           owner_email
        }

        doc = Document(filename)
        replace_placeholders_in_document(doc, placeholders)
        doc.save(filename)

        upload_url = (
            f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
            f"/drive/items/{sow_folder['id']}:/{filename}:/content"
        )
        with open(filename, "rb") as fd:
            requests.put(upload_url, headers=HEADERS_MS, data=fd)

        update_sow_status(deal_id)
        print(f"✅ SOW '{filename}' uploaded for {company_name}!")

# Run SOW generation
deals_for_sow = fetch_deals_for_sow()
for deal in deals_for_sow:
    generate_sow_for_deal(deal)
print("✅ All SOWs processed!")

# ─────────────────────────────────────────────────────────────────────────────
# MSA GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def fetch_companies_for_msa():
    """
    Fetch HubSpot companies with msa_status.
    """
    url = (
        "https://api.hubapi.com/crm/v3/objects/companies?properties="
        "name,city,state_list,zip,address,msa_status"
    )
    resp = requests.get(url, headers=HEADERS_HS)
    return resp.json().get("results", []) if resp.status_code == 200 else []

def fetch_primary_contact_for_msa(company_id):
    """
    Fetch primary contact properties for MSA (firstname, lastname, email, jobtitle, msa_status).
    """
    url = f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}/associations/contacts"
    resp = requests.get(url, headers=HEADERS_HS)
    results = resp.json().get("results", []) if resp.status_code == 200 else []
    if not results:
        return {}
    cid = results[0]["id"]
    contact_url = (
        f"https://api.hubapi.com/crm/v3/objects/contacts/{cid}?properties="
        "firstname,lastname,email,jobtitle,msa_status"
    )
    contact_resp = requests.get(contact_url, headers=HEADERS_HS)
    if contact_resp.status_code == 200:
        props = contact_resp.json().get("properties", {})
        props["id"] = cid
        return props
    return {}

def update_contact_msa_status(contact_id):
    """
    Update a contact's msa_status to 'Generated'.
    """
    if not contact_id:
        return
    url = f"https://api.hubapi.com/crm/v3/objects/contacts/{contact_id}"
    resp = requests.patch(
        url,
        headers=HEADERS_HS,
        json={"properties": {"msa_status": "Generated"}}
    )
    if resp.status_code != 200:
        send_error_email("Contact MSA Status Update Failed", resp.text)

def update_company_msa_status(company_id):
    """
    Update a company's msa_status to 'generated'.
    """
    if not company_id:
        return
    url = f"https://api.hubapi.com/crm/v3/objects/companies/{company_id}"
    resp = requests.patch(
        url,
        headers=HEADERS_HS,
        json={"properties": {"msa_status": "generated"}}
    )
    if resp.status_code != 200:
        send_error_email("Company MSA Status Update Failed", resp.text)

def msa_file_exists(folder_id, prefix):
    """
    Check if an MSA file already exists (prefix match) in the given folder.
    """
    url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{folder_id}/children"
    resp = requests.get(url, headers=HEADERS_MS)
    return any(f["name"].startswith(prefix) for f in resp.json().get("value", []))

def generate_msa_for_company(company):
    """
    For each company, generate and upload an MSA if needed.
    """
    company_id = company["id"]
    props = company["properties"]
    company_name = props.get("name", "Unknown Company")

    allow_subfolders = client_folders.get(company_name, {}).get("allow_subfolders", True)

    contact = fetch_primary_contact_for_msa(company_id)
    contact_status = (contact.get("msa_status") or "").lower().strip()

    if contact_status != "generate":
        return

    # Locate company folder in SharePoint
    parent_id = VENDORS_PARTNERS_FOLDER_ID if not allow_subfolders else CLIENTS_FOLDER_ID
    url_fldr = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{parent_id}/children"
    )
    folders = requests.get(url_fldr, headers=HEADERS_MS).json().get("value", [])
    company_folder = next((f for f in folders if f["name"] == company_name), None)
    if not company_folder:
        send_error_email("MSA Error", f"Folder not found for {company_name}")
        return

    # Determine target folder for MSA
    if allow_subfolders:
        # Create MSA subfolder if it doesn't exist
        msa_folder_id = get_or_create_subfolder(
            company_folder["id"],
            "05. MSAs",
            SUBFOLDER_05_MSAS_ID
        )
        if not msa_folder_id:
            send_error_email("MSA Error", f"MSA subfolder missing for {company_name}")
            return
        target_folder_id = msa_folder_id
    else:
        # Vendors/partners: use company folder directly
        target_folder_id = company_folder["id"]

    prefix = f"AMZ Risk - MSA - {company_name}"
    date_str = datetime.now().strftime('%Y%m%d')
    filename = f"{prefix} - {date_str}.docx"

    # Check if file already exists
    children_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{target_folder_id}/children"
    )
    files = requests.get(children_url, headers=HEADERS_MS).json().get("value", [])
    if any(f["name"] == filename for f in files):
        update_contact_msa_status(contact.get("id"))
        return

    # Copy template to target folder
    copy_url = f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}/drive/items/{MSA_TEMPLATE_ID}/copy"
    payload = {"parentReference": {"id": target_folder_id}, "name": filename}
    copy_resp = requests.post(copy_url, headers=HEADERS_MS, json=payload)
    if copy_resp.status_code not in [200, 202]:
        send_error_email("MSA Copy Failed", copy_resp.text)
        return
    time.sleep(5)
    files = requests.get(children_url, headers=HEADERS_MS).json().get("value", [])
    new_file = next((f for f in files if f["name"] == filename), None)
    if not new_file:
        send_error_email("MSA Missing", f"Copied MSA not found for {company_name}")
        return
    download_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{new_file['id']}/content"
    )
    file_data = requests.get(download_url, headers=HEADERS_MS).content
    with open(filename, "wb") as f:
        f.write(file_data)
    doc = Document(filename)
    replacements = {
        "{date}":       datetime.now().strftime("%Y-%m-%d"),
        "legal_entity_name}":       company_name,
        "{address}":    props.get("address", ""),
        "{city}":       props.get("city", ""),
        "{state_list}": props.get("state_list", ""),
        "{zip}":        props.get("zip", ""),
        "{email}":      contact.get("email", ""),
        "{firstname}":  contact.get("firstname", ""),
        "{lastname}":   contact.get("lastname", ""),
        "{jobtitle}":   contact.get("jobtitle", "")
    }
    replace_placeholders_in_document(doc, replacements)
    doc.save(filename)
    upload_url = (
        f"{GRAPH_API_BASE_URL}/sites/{SHAREPOINT_SITE_ID}"
        f"/drive/items/{target_folder_id}:/{filename}:/content"
    )
    with open(filename, "rb") as f:
        requests.put(upload_url, headers=HEADERS_MS, data=f)
    update_contact_msa_status(contact.get("id"))
    print(f"✅ MSA '{filename}' created and uploaded for {company_name}!")

# Run MSA generation
companies_for_msa = fetch_companies_for_msa()
for comp in companies_for_msa:
    generate_msa_for_company(comp)
print("✅ All MSAs processed!")

# ─────────────────────────────────────────────────────────────────────────────
# MAIN EXECUTION
# ─────────────────────────────────────────────────────────────────────────────

# In the main execution, ensure both NDA and MSA generation are called for each company
companies_list = fetch_all_hubspot_data("companies", COMPANY_FIELDS)

for comp in companies_list:
    generate_nda_for_company(comp)
    generate_msa_for_company(comp)
print("✅ All NDAs and MSAs processed!")

# Sync Closed-Won deals to Asana
sync_closed_won_deals_to_asana()
